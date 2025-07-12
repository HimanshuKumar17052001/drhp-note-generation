from fastapi import APIRouter, HTTPException, Header, Request, Depends
from fastapi.responses import FileResponse
from fastapi.background import BackgroundTasks
from app.services.auth_utils import login_required
from bson import ObjectId
from math import isnan
import tempfile
import os
from dotenv import load_dotenv
from app.models.schemas import User, BseChecklist, SebiChecklist, StandardChecklist, Company
from app.utils.log_to_s3 import log_to_s3
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

load_dotenv()
#new changesssjsjjsjsj /home/ubuntu/drhp-analyser-new/DRHP_crud_backend/app/api/endpoints/drhp_report.py
router = APIRouter()

@router.get("/drhp/report/{company_id}/download")
@log_to_s3
async def download_drhp_report(
    company_id: str,
    background_tasks: BackgroundTasks,
    authorization: str = Header(...),
    current_user: User = Depends(login_required)
):
    try:
        if not ObjectId.is_valid(company_id):
            raise HTTPException(status_code=400, detail="Invalid company ID format")

        # Helper function to safely get values
        def safe_get(value):
            if isinstance(value, float) and isnan(value):
                return ''
            if isinstance(value, list):
                # If it's a list, join the elements or take the first element
                if len(value) > 0:
                    return ' '.join(str(item).strip() for item in value if item)
                else:
                    return ''
            return value if value is not None else ''

        # Collect all flagged items from all three checklists
        all_flagged_items = []

        # 1. Load BSE checklist and filter for flagged status
        bse_items = BseChecklist.objects(company_id=ObjectId(company_id), status="FLAGGED")
        for item in bse_items:
            all_flagged_items.append({
                "summary_analysis": safe_get(item.summary_analysis)
            })

        # 2. Load SEBI checklist and filter for flagged status
        sebi_items = SebiChecklist.objects(company_id=ObjectId(company_id), status="FLAGGED")
        for item in sebi_items:
            all_flagged_items.append({
                "summary_analysis": safe_get(item.summary_analysis)
            })

        # 3. Load Standard checklist and filter for flagged status
        standard_items = StandardChecklist.objects(company_id=ObjectId(company_id), status="FLAGGED")
        for item in standard_items:
            all_flagged_items.append({
                "summary_analysis": safe_get(item.summary_analysis)
            })
            
        company_name = Company.objects(id=ObjectId(company_id)).first().name
        
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('DRHP Report - Flagged Items', 0)
        title.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add company info
        doc.add_paragraph(f'Company Name: {company_name}')
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total flagged items: {len(all_flagged_items)}')
        doc.add_paragraph('')

        if all_flagged_items:
            # Create table with 2 columns
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S.No'
            hdr_cells[1].text = 'Summary Analysis'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for index, item in enumerate(all_flagged_items, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(index)
                row_cells[1].text = str(item['summary_analysis'])

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(1.0)
                row.cells[1].width = Inches(5.0)
        else:
            doc.add_paragraph('No flagged items found for this company.')

        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_filename = temp_file.name
            doc.save(temp_filename)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"drhp_report_{company_name}_{timestamp}.docx"
        
        # Schedule file deletion after response
        background_tasks.add_task(os.remove, temp_filename)
        
        return FileResponse(
            path=temp_filename,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print(f"Error generating DRHP report: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e)) 